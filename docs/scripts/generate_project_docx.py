from docx import Document
from docx.shared import Pt
from datetime import date


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_par(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Titelseite
    add_par(doc, "HTL Saalfelden", bold=True)
    add_par(doc, "Systemplanung und Projektentwicklung", bold=True)
    doc.add_paragraph()
    add_heading(doc, "Projektdokumentation 2025 / 2026", level=0)

    add_table(
        doc,
        ["Feld", "Wert"],
        [
            ["Projektbezeichnung", "GOOSE"],
            ["Projektteam", "Team 8"],
            ["Erstellt am", "28.03.2026"],
            ["Letzte Änderung am", "28.03.2026"],
            ["Status", "in Bearbeitung"],
            ["Aktuelle Version", "1.0"],
        ],
    )

    doc.add_page_break()

    # Änderungsverlauf
    add_heading(doc, "Änderungsverlauf", level=1)
    add_table(
        doc,
        ["Nr.", "Datum", "Version", "Geänderte Kapitel", "Art der Änderung", "Autor"],
        [
            ["1", "26.03.2026", "1.0", "Alle", "Erstellung", "Team 8"],
            ["2", "28.03.2026", "1.0", "Kapitel 1-8", "Inhaltliche Vervollständigung", "Team 8"],
        ],
    )

    doc.add_page_break()

    # 1 Allgemeines
    add_heading(doc, "1. Allgemeines / Projektübersicht", level=1)
    add_heading(doc, "1.1 Projektbeschreibung", level=2)
    add_par(
        doc,
        "GOOSE ist ein Echtzeit-System zur autonomen und manuellen Steuerung einer DJI Tello Drohne. "
        "Das System kombiniert KI-Objekterkennung (YOLO), geometrische Positionsschätzung (Ellipse/Ring), "
        "Kalman-Filterung sowie einen zustandsbasierten Autopilot (ALIGN, APPROACH, PUNCH). "
        "Zusätzlich unterstützt GOOSE Datenerfassung für Retraining, Logging und Bedienung via Tastatur/Joystick.",
    )

    add_heading(doc, "1.2 Projektteam und Schnittstellen", level=2)
    add_table(
        doc,
        ["Rolle", "Name", "Telefon", "E-Mail", "Team"],
        [
            ["Member", "Emilio Schwaiger", "Zu Privat", "QuackingQuark@gmail.com", "Team 8"],
            ["Member", "Leonhard Fresacher", "Zu Privat", "ohSystemmm@gmail.com", "Team 8"],
            ["Member", "Christoph Fercher", "Zu Privat", "chr.fercher@gmail.com", "Team 8"],
        ],
    )
    add_par(doc, "Schnittstellen: DJI Tello (UDP), Kamera-Stream, Joystick (Pygame), Roboflow API, lokales Dateisystem.")

    # 2 Funktionale Anforderungen
    add_heading(doc, "2. Funktionale Anforderungen", level=1)
    add_heading(doc, "2.1 Use Cases", level=2)

    add_heading(doc, "2.1.1 Use Case: Manueller Flug", level=3)
    add_par(doc, "Akteur: Pilot. Ziel: Sichere manuelle Steuerung inklusive Takeoff/Land/Not-Aus.")
    add_par(doc, "Ablauf: Start des Systems, Verbindungsaufbau, Eingabe via Tastatur oder Joystick, RC-Kommandos in Echtzeit.")
    add_par(doc, "Ergebnis: Drohne reagiert innerhalb der Schleifenfrequenz und bleibt steuerbar.")

    add_heading(doc, "2.1.2 Use Case: Autonomes Zentrieren und Durchflug", level=3)
    add_par(doc, "Akteur: Pilot (Autopilot aktivieren). Ziel: Automatisches Ausrichten, Annähern und Durchflug durch Zielring.")
    add_par(doc, "Ablauf: Vision erkennt Ziel, PositionEstimator berechnet Pose, AutoPilot steuert Phasen ALIGN/APPROACH/PUNCH.")
    add_par(doc, "Ergebnis: Stabiler, reproduzierbarer Anflug mit manueller Übersteuerbarkeit.")

    add_heading(doc, "2.1.3 Use Case: Datenerfassung für Modellverbesserung", level=3)
    add_par(doc, "Akteur: Entwicklerteam. Ziel: Speichern von Kontext-Frames und unsicheren Samples zur Label-Nacharbeit.")
    add_par(doc, "Ablauf: Während Flug werden Frames und Metadaten in flight_data abgelegt; Upload-Skripte übertragen Daten zu Roboflow.")
    add_par(doc, "Ergebnis: Kontinuierliche Verbesserung der Erkennungsqualität.")

    # 3 Nichtfunktional
    add_heading(doc, "3. Nichtfunktionale Anforderungen", level=1)
    add_table(
        doc,
        ["Kategorie", "Anforderung"],
        [
            ["Echtzeit", "Bildverarbeitung und Steuerung laufen mit geringer Latenz (Zielbereich ~30 FPS)."],
            ["Zuverlässigkeit", "Ausfallsichere Fallbacks (Coasting, Schwellenwerte, manuelle Übernahme)."],
            ["Sicherheit", "Not-Aus, manuelle Priorisierung, Begrenzung kritischer Flugkommandos."],
            ["Wartbarkeit", "Modulare Architektur (core/vision), klar getrennte Verantwortlichkeiten."],
            ["Portabilität", "Betrieb auf Windows mit Python-Umgebung und reproduzierbarer requirements.txt."],
            ["Nachvollziehbarkeit", "Rotierende Logfiles, strukturierte Metadaten bei Datenerfassung."],
        ],
    )

    # 4 Projektplanung
    add_heading(doc, "4. Projektplanung", level=1)
    add_heading(doc, "4.1 Variantenbildung", level=2)
    add_par(doc, "Untersuchte Varianten:")
    add_par(doc, "Variante A: Reine ArUco-basierte Führung. Vorteil: präzise Marker-Lokalisierung; Nachteil: starke Szenenabhängigkeit.")
    add_par(doc, "Variante B: Reine KI-Erkennung (YOLO). Vorteil: flexibel; Nachteil: Tiefenschätzung ohne Geometrie weniger stabil.")
    add_par(doc, "Variante C (gewählt): Hybrider Ansatz aus YOLO + geometrischer Pose + Kalman-Filter.")

    add_heading(doc, "4.2 Machbarkeitsstudie", level=2)
    add_par(doc, "Vorab validiert wurden: Tello-Verbindung über djitellopy, Echtzeit-Frame-Verarbeitung mit OpenCV, YOLO-Inferenz (.pt/.onnx), "
                "Kalman-basierte Glättung und RC-Steuerung im Live-Loop.")

    add_heading(doc, "4.3 Allgemeine Planungsinformationen", level=2)
    add_par(doc, "Programmiersprache: Python 3.8+. Betriebssystem: Windows. Projektverwaltung über Git. "
                "Iterationsbasiertes Vorgehen mit kurzen Validierungszyklen auf echter Hardware.")

    add_heading(doc, "4.4 Projektumfeldanalyse", level=2)
    add_par(doc, "Vergleichbare Lösungen: akademische UAV-Demos, OpenCV-/ArUco-Drohnenprojekte und YOLO-Tracking-Prototypen.")
    add_par(doc, "Abgrenzung: Fokus auf robusten, praxisnahen Tello-Workflow mit kombinierter Erkennung, geometrischer Pose, "
                "Autopilot-Phasen und integrierter Datenerfassung.")
    add_par(doc, "Stakeholder: Team 8 (Entwicklung), Lehrende/Bewertungsgremium, Demonstrationspublikum, potenzielle Entwickler-Nutzer.")

    # 5 Softwarearchitektur
    add_heading(doc, "5. Softwarearchitektur", level=1)
    add_par(doc, "Das System besteht aus einer Hauptschleife (main.py), einer Vision-Worker-Thread-Pipeline und modularen Kernkomponenten.")

    add_heading(doc, "5.1 Aktivitätsdiagramme", level=2)
    add_heading(doc, "5.1.1 Aktivitätsdiagramm: Flugzyklus", level=3)
    add_par(doc, "Aktivität: Initialisierung -> Frame holen -> Detektion -> Pose -> Steuerentscheidung -> RC senden -> Rendering -> Wiederholen.")
    add_heading(doc, "5.1.2 Aktivitätsdiagramm: Datenerfassung", level=3)
    add_par(doc, "Aktivität: Trigger Datenerfassung -> Sample-Auswahl -> Speichern Bild + JSONL-Metadaten -> optional Upload.")

    add_heading(doc, "5.2 Sequenzdiagramme", level=2)
    add_heading(doc, "5.2.1 Sequenzdiagramm: Autopilot", level=3)
    add_par(doc, "main.py fordert Pose von VisionWorker an; AutoPilot berechnet RC-Werte; DroneController sendet UDP-Kommandos an Tello.")
    add_heading(doc, "5.2.2 Sequenzdiagramm: Uncertain Upload", level=3)
    add_par(doc, "Sampler markiert unsichere Frames; upload_uncertain_data.py liest Dateien; Roboflow API nimmt Upload entgegen.")

    add_heading(doc, "5.3 Komponentendiagramme", level=2)
    add_par(doc, "Vorhanden in GOOSE/diagrams/componentdiagram.wsd. Hauptkomponenten: DroneController, VisionWorker, ObjectDetector, "
                "PositionEstimator, AutoPilot, UI/ControlCenter, DataFactory.")

    add_heading(doc, "5.4 Verteilungsdiagramme", level=2)
    add_par(doc, "Verteilung auf zwei physische Knoten: (1) Host-PC/Notebook führt Python-Anwendung aus, "
                "(2) DJI Tello liefert Videostream und empfängt RC-Kommandos über WLAN/UDP.")

    add_heading(doc, "5.5 Softwarekomponenten / Programme", level=2)
    add_heading(doc, "5.5.1 SW Programme", level=3)
    add_table(
        doc,
        ["Programm", "Version (Projekt)"],
        [
            ["Python", "3.8.x"],
            ["Visual Studio Code / PyCharm", "aktuelle Schul-/Projektversion"],
            ["Git", "2.x"],
        ],
    )

    add_heading(doc, "5.5.2 SW Komponenten", level=3)
    add_table(
        doc,
        ["Komponente", "Hersteller", "Lizenz", "Bezugsquelle"],
        [
            ["numpy", "NumPy Developers", "BSD", "https://pypi.org/project/numpy/"],
            ["opencv-contrib-python", "OpenCV", "Apache 2 / BSD", "https://pypi.org/project/opencv-contrib-python/"],
            ["djitellopy", "djitellopy Contributors", "MIT", "https://pypi.org/project/djitellopy/"],
            ["pygame", "pygame community", "LGPL", "https://pypi.org/project/pygame/"],
            ["kivy", "Kivy Organization", "MIT", "https://pypi.org/project/kivy/"],
            ["ultralytics", "Ultralytics", "AGPL-3.0", "https://pypi.org/project/ultralytics/"],
            ["onnxruntime", "Microsoft", "MIT", "https://pypi.org/project/onnxruntime/"],
            ["python-dotenv", "Python Dotenv Authors", "BSD", "https://pypi.org/project/python-dotenv/"],
        ],
    )

    # 6 Projektdurchführung
    add_heading(doc, "6. Projektdurchführung", level=1)

    add_heading(doc, "6.1 Sprint 1", level=2)
    add_heading(doc, "6.1.1 Sprintplanung", level=3)
    add_par(doc, "Dauer: 12.03.2026 - 18.03.2026")
    add_par(doc, "Ziele: Basisflug stabilisieren, YOLO-Detektion in Hauptloop integrieren, Logging und Aufzeichnungsstruktur aufsetzen.")
    add_par(doc, "Geplante Story Points: 20")
    add_heading(doc, "6.1.2 Sprint Demo", level=3)
    add_par(doc, "Erfolgreich demonstriert: stabiler Stream, manuelle Steuerung, erste zielbasierte Detektion mit Bounding Boxes.")
    add_heading(doc, "6.1.3 Sprint Retrospektive", level=3)
    add_par(doc, "Positiv: schnelle Integrationsfortschritte. Herausforderungen: schwankende Erkennung bei schwieriger Beleuchtung.")
    add_heading(doc, "6.1.4 Sprint Zusammenfassung", level=3)
    add_par(doc, "Ergebnis: lauffähiger End-to-End-Loop als Grundlage für Autopilot und Pose-Feinschliff.")

    add_heading(doc, "6.2 Sprint 2", level=2)
    add_heading(doc, "6.2.1 Sprintplanung", level=3)
    add_par(doc, "Dauer: 19.03.2026 - 25.03.2026")
    add_par(doc, "Ziele: Pose-Schätzung verbessern, Autopilot-Phasen ausarbeiten, Datenpipeline (context/uncertain) etablieren.")
    add_par(doc, "Geplante Story Points: 24")
    add_heading(doc, "6.2.2 Sprint Demo", level=3)
    add_par(doc, "Erfolgreich demonstriert: stabile Zielwahl, Kalman-Coasting, Phasensteuerung ALIGN/APPROACH/PUNCH, Datenspeicherung in JSONL.")
    add_heading(doc, "6.2.3 Sprint Retrospektive", level=3)
    add_par(doc, "Positiv: bessere Regelstabilität. Offen: weitere Tuningarbeit bei starkem Bildrauschen und Randfällen.")
    add_heading(doc, "6.2.4 Sprint Zusammenfassung", level=3)
    add_par(doc, "Ergebnis: deutlich robustere autonome Fluglogik mit messbarer Datenbasis für weiteres Training.")

    add_heading(doc, "6.3 Sprint n", level=2)
    add_par(doc, "Nächste Schwerpunkte: Testsystematik erweitern, UI/ControlCenter ausbauen, Deployment-Dokumentation finalisieren.")

    # 7 Installation
    add_heading(doc, "7. Installation / Software deployment", level=1)
    add_par(doc, "1) Python 3.8+ installieren und virtuelle Umgebung erstellen.")
    add_par(doc, "2) Abhängigkeiten installieren: pip install -r requirements.txt")
    add_par(doc, "3) Modelle im Ordner GOOSE/assets/models bereitstellen (targetModel.onnx oder targetModel.pt).")
    add_par(doc, "4) Mit Tello-WLAN verbinden.")
    add_par(doc, "5) Start: python GOOSE/main.py --model onnx --ip 192.168.10.1")
    add_par(doc, "6) Optional: Joystick-Konfiguration unter GOOSE/flight_data/joystick_config.json prüfen.")
    add_par(doc, "7) Optional: Datensammlung aktivieren und Ergebnisse unter flight_data prüfen.")

    # 8 Abschluss
    add_heading(doc, "8. Projektabschluß", level=1)
    add_heading(doc, "8.1 Projektzusammenfassung", level=2)
    add_par(doc, "GOOSE erreicht das Kernziel eines robusten, echtzeitfähigen Drohnen-Demonstrators mit KI-gestützter Zielerfassung "
                "und autonomer Flugunterstützung. Besonders erfolgreich war die Kombination aus Detektion, geometrischer Pose und "
                "Filterung. Verbesserungspotenzial besteht in automatisierten Tests, weiterem Hyperparameter-Tuning und noch stärkerer "
                "Standardisierung der Deployment-Pfade.")

    add_heading(doc, "8.2 Attachments", level=2)
    add_table(
        doc,
        ["Datei", "Beschreibung"],
        [
            ["GOOSE/main.py", "Haupteinstieg und Event-/Control-Loop"],
            ["GOOSE/core/autopilot.py", "Autopilot-Zustandsautomat und Steuerlogik"],
            ["GOOSE/vision/position_estimator.py", "3D-Positionsschätzung mit Kalman-Filter"],
            ["GOOSE/vision/detector.py", "YOLO-Detektionsadapter"],
            ["GOOSE/core/drone.py", "Drohnenkommunikation (djitellopy)"],
            ["GOOSE/diagrams/componentdiagram.wsd", "Komponentendiagramm"],
            ["requirements.txt", "Python-Abhängigkeiten"],
            ["PROGRAM_DOCUMENTATION.md", "Projektweite technische Kurzdokumentation"],
        ],
    )

    out_path = "GOOSE_Projektdokumentation_Team8_v1.0.docx"
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    build()
